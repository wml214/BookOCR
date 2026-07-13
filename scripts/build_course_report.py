"""生成课程论文格式的图书盘点系统报告初稿。

说明：
1. 本脚本使用 python-docx 生成 DOCX 报告，便于后续按《电子学报》模板继续调整。
2. 由于教师给出的模板是旧版 .doc，本脚本不直接修改模板文件，而是生成一份内容完整、
   结构清晰的小论文初稿。
3. 文中姓名、学号、小组成员等信息使用占位符，提交前需要人工替换。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOCX = PROJECT_ROOT / "docs/paper/基于视觉的图书盘点系统课程报告_初稿.docx"


def set_run_font(run, font_name: str = "宋体", size: float | None = None, bold: bool | None = None) -> None:
    """设置中英文字体，避免 Word 打开后中文回退到随机字体。"""

    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 4, line: float = 1.15) -> None:
    """设置段落前后距和行距。"""

    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_paragraph(doc: Document, text: str = "", *, style: str | None = None, first_line: bool = True):
    """添加正文段落，并统一套用中文小论文正文格式。"""

    paragraph = doc.add_paragraph(style=style)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    set_paragraph_spacing(paragraph, after=4, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 10.5)
    return paragraph


def add_heading(doc: Document, title: str, level: int = 1) -> None:
    """添加章节标题。"""

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(title)
    set_run_font(run, "黑体", 12 if level == 1 else 11, True)


def set_cell_text(cell, text: str, *, bold: bool = False, align_center: bool = False) -> None:
    """设置表格单元格文本格式。"""

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 9, bold)


def set_table_borders(table) -> None:
    """给表格添加细边框。"""

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "808080")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    """添加带标题的三线风格数据表。"""

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap, before=4, after=2)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, "黑体", 9.5, True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, align_center=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, align_center=index > 0)

    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str, width_cm: float = 12.0) -> None:
    """插入图片和图题；图片不存在时自动跳过。"""

    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap, after=4)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, "黑体", 9.5, False)


def add_formula(doc: Document, text: str) -> None:
    """以居中形式添加评价指标公式。"""

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=2, after=2)
    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman", 10.5)


def configure_document(doc: Document) -> None:
    """设置页面、默认字体和页脚。"""

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("基于视觉的图书盘点系统课程报告")
    set_run_font(run, "宋体", 9)


def build_report() -> None:
    """生成课程论文初稿。"""

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=6, after=8)
    title_run = title.add_run("基于视觉的图书盘点系统设计与实现")
    set_run_font(title_run, "黑体", 18, True)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(authors, after=3)
    author_run = authors.add_run("学生：姓名1，姓名2，姓名3，姓名4    学号：XXXXXXXX    指导教师：XXX")
    set_run_font(author_run, "宋体", 10.5)

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(org, after=8)
    org_run = org.add_run("（学院/专业/班级，天津，300000）")
    set_run_font(org_run, "宋体", 10)

    abstract = doc.add_paragraph()
    set_paragraph_spacing(abstract, after=5, line=1.15)
    run = abstract.add_run("摘  要：")
    set_run_font(run, "黑体", 10.5, True)
    run = abstract.add_run(
        "针对高校图书馆开放书架盘点中人工记录效率低、重复统计和漏检较难避免的问题，"
        "本文设计并实现了一套基于视觉的图书盘点系统。系统以多张有序书架图像为输入，"
        "首先采用 YOLOv8n-OBB 模型定位倾斜书脊，再通过透视矫正获得单本书脊裁剪图，"
        "随后使用 PaddleOCR 识别书脊文字，并结合馆藏目录进行规范书名匹配。针对相邻照片"
        "存在重叠区域的情况，系统利用 ORB 特征与 RANSAC 单应性矩阵完成图像配准，并依据"
        "空间重合度与书名一致性进行重复合并。实验表明，最终 OBB 书脊检测模型在 20 张独立"
        "测试图像、456 个书脊实例上取得 Precision=0.966、Recall=0.941、mAP@0.5=0.988 的结果；"
        "在推荐阈值下预测书脊数为 452，书脊数量计数准确率为 99.12%；在已检测书脊口径下，"
        "书名自动识别准确率为 96.02%。系统基于 Streamlit 实现了多图上传、结果预览、人工复核"
        "和 CSV 导出功能，能够完成从图像采集到盘点表生成的完整流程。"
    )
    set_run_font(run, "宋体", 10.5)

    keywords = doc.add_paragraph()
    set_paragraph_spacing(keywords, after=8)
    run = keywords.add_run("关键词：")
    set_run_font(run, "黑体", 10.5, True)
    run = keywords.add_run("图书盘点；书脊检测；YOLOv8-OBB；PaddleOCR；馆藏匹配；Streamlit")
    set_run_font(run, "宋体", 10.5)

    english_title = doc.add_paragraph()
    english_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(english_title, before=4, after=4)
    run = english_title.add_run("Design and Implementation of a Vision-based Book Inventory System")
    set_run_font(run, "Times New Roman", 13, True)

    english_abs = doc.add_paragraph()
    set_paragraph_spacing(english_abs, after=5)
    run = english_abs.add_run("Abstract: ")
    set_run_font(run, "Times New Roman", 10.5, True)
    run = english_abs.add_run(
        "A vision-based book inventory system is proposed for bookshelf inspection. "
        "The system detects book spines with YOLOv8n-OBB, performs perspective correction "
        "and OCR recognition, and matches OCR text with a library catalog. Adjacent images "
        "are registered by ORB features and homography estimation to reduce duplicate counting. "
        "Experiments on an independent test set show that the proposed detector achieves "
        "0.966 precision, 0.941 recall and 0.988 mAP@0.5. The book-spine count accuracy reaches "
        "99.12%, and the title recognition accuracy on detected spines reaches 96.02%."
    )
    set_run_font(run, "Times New Roman", 10.5)

    english_keywords = doc.add_paragraph()
    set_paragraph_spacing(english_keywords, after=8)
    run = english_keywords.add_run("Key words: ")
    set_run_font(run, "Times New Roman", 10.5, True)
    run = english_keywords.add_run("book inventory; book spine detection; YOLOv8-OBB; PaddleOCR; catalog matching")
    set_run_font(run, "Times New Roman", 10.5)

    add_heading(doc, "1 引言")
    add_paragraph(
        doc,
        "图书馆书架盘点是馆藏管理中的基础工作，传统方式通常依赖人工逐本核对书名、索书号和数量。"
        "在开放书架和密集排架场景中，图书摆放存在倾斜、遮挡、反光、文字竖排等情况，人工盘点不仅耗时，"
        "还容易出现漏记、重复记录和同名书混淆。随着目标检测、OCR 与轻量化应用框架的发展，利用普通手机"
        "拍摄书架图像并自动完成图书盘点具有较高的实践价值。"
    )
    add_paragraph(
        doc,
        "本课程设计面向西院图书馆局部书架场景，构建了不少于 200 张独立书架图像的数据集，完成书脊标注、"
        "模型训练、OCR 识别、馆藏目录匹配、跨图去重和可视化软件开发。与普通水平矩形框检测相比，书架中的"
        "书脊经常倾斜放置，本文最终采用旋转目标框（Oriented Bounding Box, OBB）进行书脊定位，以提高裁剪"
        "方向和后续 OCR 的稳定性。"
    )

    add_heading(doc, "2 系统总体设计")
    add_paragraph(
        doc,
        "系统整体采用“多张图像盘点”路线。用户按从左到右、从上到下的顺序拍摄相邻书架图像，并尽量保留"
        "20%~30% 重叠区域。系统首先对输入图像进行书脊检测，得到每本书的旋转框；然后根据 OBB 四点坐标"
        "进行透视矫正和方向统一，生成单本书脊裁剪图；之后调用 PaddleOCR 获得文字候选，并将 OCR 文本与"
        "馆藏目录中的规范题名、作者、索书号等字段进行相似度匹配；最后对多张相邻图像执行空间去重，输出"
        "按书名汇总后的盘点结果。"
    )

    add_table(
        doc,
        ["模块", "输入", "输出", "作用"],
        [
            ["书脊检测", "书架图像", "书脊 OBB 坐标和置信度", "定位每一本实体书的书脊区域"],
            ["书脊裁剪", "OBB 四点坐标", "矫正后的单本书脊图", "改善倾斜和竖排文字的 OCR 条件"],
            ["OCR 识别", "书脊裁剪图", "OCR 原文和置信度", "提取书名、作者、索书号等文本线索"],
            ["馆藏匹配", "OCR 文本、馆藏目录", "规范书名和匹配分数", "将不完整 OCR 结果纠正为馆藏题名"],
            ["跨图去重", "相邻图像、OBB 坐标、匹配书名", "去重后的盘点结果", "避免重叠区域重复计数"],
            ["软件界面", "用户上传图片和参数", "结果图、CSV、日志", "提供完整演示和人工复核入口"],
        ],
        "表1 系统主要模块及功能",
    )

    add_heading(doc, "3 数据集构建与标注")
    add_paragraph(
        doc,
        "原始数据由 102 张 HEIC 图像和后续补拍的 100 张 JPG 图像组成。HEIC 图像经 pillow-heif 转换为 JPEG，"
        "补拍图像统一整理到中间数据目录，最终形成 202 张独立书架图像。压缩包与原 HEIC 图像属于同一批素材，"
        "未重复计入数据量。馆藏目录来自西院图书馆 Excel 清单，经预处理后得到包含约 59078 条规范书目记录的"
        "CSV 文件，用于后续书名匹配。"
    )
    add_paragraph(
        doc,
        "书脊标注使用 Label Studio 完成。考虑到大量书籍存在倾斜摆放，标注时以多边形轮廓覆盖每本书的可见"
        "书脊区域，并在导出阶段转换为 YOLOv8-OBB 所需的旋转框格式。最终人工确认 200 张图像，共 4537 个书脊"
        "旋转框实例，并按训练集、验证集、测试集划分为 160/20/20 张图像。划分时尽量避免同一连续书架图像跨集合，"
        "以降低数据泄漏风险。"
    )
    add_table(
        doc,
        ["项目", "数量或说明"],
        [
            ["整理后图像总数", "202 张"],
            ["用于最终 OBB 数据集的图像", "200 张"],
            ["书脊 OBB 标注实例", "4537 个"],
            ["训练/验证/测试划分", "160 / 20 / 20 张"],
            ["馆藏目录规模", "约 59078 条规范书目"],
            ["模型输入尺寸", "640 像素"],
        ],
        "表2 数据集与馆藏目录概况",
    )

    add_heading(doc, "4 书脊检测模型")
    add_heading(doc, "4.1 模型选择", level=2)
    add_paragraph(
        doc,
        "课程设计初期曾以 YOLOv8n-seg 作为书脊分割基线模型。实验过程中发现，书架图像中存在大量倾斜书脊，"
        "普通水平矩形框或分割轮廓在裁剪时容易引入背景和相邻书脊，影响 OCR 识别。为此，本文最终选择"
        "YOLOv8n-OBB 作为主模型。OBB 模型直接输出旋转矩形框，能够更准确描述倾斜书脊方向，便于后续进行"
        "透视矫正和竖排文字识别。"
    )
    add_heading(doc, "4.2 训练设置", level=2)
    add_paragraph(
        doc,
        "模型基于 Ultralytics YOLOv8n-OBB 预训练权重进行迁移学习，训练环境为 Windows、Python 3.10、"
        "PyTorch CUDA 版本和 NVIDIA RTX 4060 8GB 显卡。训练过程中使用单类别 book_spine，输入尺寸为 640，"
        "最终模型权重保存为 models/weights/book_spine_obb_final_200_best.pt。推理阶段采用 conf=0.60、"
        "iou=0.50 作为默认阈值，以减少重复框并保持较高召回率。"
    )
    add_figure(
        doc,
        PROJECT_ROOT / "outputs/reports/book_spine_obb_final_200_test/val_batch0_pred.jpg",
        "图1 测试集书脊 OBB 检测结果示例",
        width_cm=13.5,
    )

    add_heading(doc, "5 OCR 识别与馆藏目录匹配")
    add_paragraph(
        doc,
        "检测得到书脊 OBB 后，系统按照四点坐标对书脊区域进行透视变换。裁剪模块会根据书脊长短边判断方向，"
        "将书脊统一为便于 OCR 处理的纵向或横向图像，并保留少量边缘 padding，避免文字贴边导致识别不完整。"
        "OCR 部分采用 PaddleOCR，输出每个书脊裁剪图中的文本行、平均置信度和合并后的 OCR 原文。"
    )
    add_paragraph(
        doc,
        "由于书脊文字存在竖排、反光、遮挡和字体较小等问题，单纯依赖 OCR 字符串很难保证书名完全正确。本文"
        "引入馆藏目录作为先验知识，对 OCR 文本进行规范化处理，包括去除空格和标点、统一大小写、提取可能的"
        "索书号片段和作者片段等。匹配阶段先通过字符片段倒排索引召回候选书目，再综合完整串相似度、最长公共"
        "片段覆盖率、字符 Jaccard 相似度、作者线索和索书号线索得到匹配分数。分数达到阈值的结果标记为 matched，"
        "证据不足的结果保留为 pending，交由界面人工复核。"
    )

    add_heading(doc, "6 多图去重与计数方法")
    add_paragraph(
        doc,
        "多张有序照片进行盘点时，相邻图像的重叠区域可能包含同一本实体书。若直接汇总每张图的检测结果，容易"
        "产生重复计数。本文采用“图像配准 + 空间重合 + 书名一致”的安全去重策略。系统首先对相邻图像提取 ORB"
        "特征，使用 KNN 匹配和 Lowe ratio 筛选有效匹配点，再通过 RANSAC 估计从后一张图到前一张图的单应性矩阵。"
        "若内点数量和内点比例不足，则认为配准不可靠，不强制合并。"
    )
    add_paragraph(
        doc,
        "当配准可靠时，系统将后一张图中的书脊 OBB 映射到前一张图坐标系，计算映射后 OBB 与前一张图中 OBB 的"
        "空间 IoU。只有当两个书脊均已匹配到同一规范书名，且空间 IoU 超过阈值时，才合并为同一本实体书。该策略"
        "避免了仅按书名合并导致的错误，例如两本同名实体书并排摆放时仍应计为两册。"
    )
    add_formula(doc, "识别准确率 = 正确识别的图书本数 / 测试集中图书总数量")
    add_formula(doc, "书脊数量计数准确率 = 模型预测书脊数 / 测试集真实书脊数")
    add_formula(doc, "Precision = TP / (TP + FP)，Recall = TP / (TP + FN)")

    add_heading(doc, "7 软件系统设计与实现")
    add_paragraph(
        doc,
        "软件采用 Streamlit 实现，便于在普通电脑上快速演示。界面支持一次上传多张 JPG、PNG、HEIC 或 HEIF 图像，"
        "侧边栏可配置模型权重、馆藏目录、检测阈值、匹配阈值和去重参数。用户点击“开始盘点”后，系统依次执行"
        "图片保存、模型加载、书脊检测、OCR、馆藏匹配和可选跨图去重，并在页面上展示汇总结果、识别明细、标注"
        "预览图和运行日志。"
    )
    add_paragraph(
        doc,
        "软件输出包括盘点明细 CSV、按书名汇总的盘点结果 CSV、带 OBB 框的预览图、书脊裁剪图，以及启用去重时"
        "生成的配准日志和重复合并记录。对于 pending 结果，系统不强制错误匹配，而是保留 OCR 原文和裁剪图，"
        "便于后续人工确认。"
    )
    add_figure(
        doc,
        PROJECT_ROOT / "outputs/inventory/streamlit_20260705_141826/previews/001_IMG_7917_obb_preview.jpg",
        "图2 Streamlit 运行输出的书脊检测预览图",
        width_cm=13.5,
    )

    add_heading(doc, "8 实验结果与分析")
    add_heading(doc, "8.1 书脊检测实验", level=2)
    add_paragraph(
        doc,
        "最终模型在独立测试集上进行评估。测试集包含 20 张图像、456 个书脊实例。实验结果表明，OBB 模型能够"
        "较好适应倾斜、密集和低对比度书脊场景，Precision、Recall 和 mAP 指标均达到课程设计目标。"
    )
    add_table(
        doc,
        ["指标", "数值"],
        [
            ["Precision", "0.966"],
            ["Recall", "0.941"],
            ["F1", "0.953"],
            ["mAP@0.5", "0.988"],
            ["mAP@0.5:0.95", "0.826"],
        ],
        "表3 最终 OBB 模型测试集检测指标",
    )
    add_figure(
        doc,
        PROJECT_ROOT / "outputs/reports/book_spine_obb_final_200_test/BoxPR_curve.png",
        "图3 OBB 模型 Precision-Recall 曲线",
        width_cm=11.5,
    )
    add_figure(
        doc,
        PROJECT_ROOT / "outputs/reports/book_spine_obb_final_200_test/confusion_matrix_normalized.png",
        "图4 书脊检测归一化混淆矩阵",
        width_cm=10.5,
    )

    add_heading(doc, "8.2 书名识别与计数实验", level=2)
    add_paragraph(
        doc,
        "在推荐推理阈值 conf=0.60、iou=0.50 下，模型在测试集上预测出 452 个书脊，测试集真实书脊数为 456，"
        "书脊数量计数准确率为 452/456=99.12%。在书名识别方面，系统为 452 个检测书脊生成裁剪图并进行 OCR 与"
        "馆藏匹配，其中 434 个结果自动标记为 matched，18 个结果保留为 pending。经人工核验，matched 行未发现"
        "错误匹配，因此以已检测书脊为分母的书名自动识别准确率为 434/452=96.02%；若以测试集真实书脊数为分母，"
        "端到端识别准确率为 434/456=95.18%。"
    )
    add_table(
        doc,
        ["项目", "结果"],
        [
            ["测试集图像数", "20 张"],
            ["测试集真实书脊数", "456 个"],
            ["系统检测书脊数", "452 个"],
            ["书脊数量计数准确率", "452 / 456 = 99.12%"],
            ["自动匹配 matched 数", "434 个"],
            ["待确认 pending 数", "18 个"],
            ["已检测书脊口径书名识别准确率", "434 / 452 = 96.02%"],
            ["端到端书名识别准确率", "434 / 456 = 95.18%"],
        ],
        "表4 书名识别与计数实验结果",
    )
    add_paragraph(
        doc,
        "从错误来源看，漏检主要发生在边缘不完整书脊、低对比度书脊或被遮挡书脊处；pending 样本主要由 OCR 未能"
        "识别出完整题名、只识别出作者或索书号片段导致。由于系统将低置信度结果保留为待确认，避免了把错误书名"
        "强行计入自动匹配结果，因此自动 matched 结果具有较高可靠性。"
    )

    add_heading(doc, "8.3 消融与方案比较", level=2)
    add_paragraph(
        doc,
        "项目早期尝试使用 YOLOv8n-seg 进行书脊轮廓分割。虽然分割模型在部分样本上能够给出较好的轮廓，但在"
        "倾斜书脊和密集排列场景中，后处理裁剪较复杂，且预标注结果修改成本较高。OBB 模型直接输出旋转框，"
        "更符合书脊近似长矩形的几何特征，人工修正成本和 OCR 裁剪成本均更低。因此最终系统以 YOLOv8n-OBB 为主线。"
    )

    add_heading(doc, "9 总结与展望")
    add_paragraph(
        doc,
        "本文完成了一套基于视觉的图书盘点系统，从数据采集、书脊标注、OBB 模型训练、OCR 识别、馆藏目录匹配、"
        "跨图去重到 Streamlit 软件展示形成了完整闭环。实验结果表明，系统在独立测试集上具有较高的书脊检测精度、"
        "计数准确率和书名识别准确率，能够满足课程设计对模型训练、软件展示和评价指标的基本要求。"
    )
    add_paragraph(
        doc,
        "后续工作可从三方面改进：第一，继续补充不同书架、不同光照和更多竖排书名样本，提高模型泛化能力；第二，"
        "对 pending 样本建立人工复核反馈机制，将修正结果回流到匹配规则和 OCR 后处理；第三，在多图去重方面引入"
        "更稳定的拍摄顺序记录和全局拼接策略，提高大范围书架盘点时的重复合并稳定性。"
    )

    add_heading(doc, "参考文献")
    references = [
        "[1] Jocher G, Chaurasia A, Qiu J. Ultralytics YOLOv8. 2023.",
        "[2] PaddlePaddle Authors. PaddleOCR: Awesome multilingual OCR toolkits based on PaddlePaddle.",
        "[3] Lowe D G. Distinctive image features from scale-invariant keypoints[J]. International Journal of Computer Vision, 2004.",
        "[4] Fischler M A, Bolles R C. Random sample consensus: a paradigm for model fitting with applications to image analysis and automated cartography[J]. Communications of the ACM, 1981.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=2)
        run = p.add_run(ref)
        set_run_font(run, "宋体", 9.5)

    add_heading(doc, "附录：课程设计时间进度安排")
    add_table(
        doc,
        ["阶段", "课时", "主要内容"],
        [
            ["数据整理、标注和质量复核", "12", "HEIC 转换、图像审计、Label Studio 标注、训练/验证/测试划分"],
            ["环境搭建、工程骨架和界面原型", "4", "Conda 环境、依赖安装、项目目录、Streamlit 初始界面"],
            ["模型训练、OCR、匹配与去重集成", "16", "YOLOv8n-OBB 训练、PaddleOCR、馆藏匹配、跨图去重"],
            ["论文撰写和实验结果整理", "4", "指标统计、结果图整理、小论文撰写"],
            ["答辩演示和最终打包", "4", "演示流程、问题准备、工程压缩包整理"],
        ],
        "表5 课程设计时间进度安排",
    )

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_report()
